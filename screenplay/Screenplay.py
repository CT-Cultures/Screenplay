# -*- coding: utf-8 -*-
"""
Created on Mon Feb 17 00:31:27 2020

@author: VXhpUS
"""
from pathlib import Path
import numpy as np
import pandas as pd
import re
from docx import Document
import time
from xml.dom import minidom
import lxml
import lxml.etree as et
import codecs
import math

## For Translate
import hashlib
from random import randint
import http.client
import urllib
import json

#%% Class Screenplay
class Screenplay(object):
    
    def __init__(self, sc: str = None):   

        super(Screenplay, self).__init__()
        self.sc = sc
        self.elements = Elements()
        self.export = Export().to_OpenXML()
        self.groupSH2Name = {
                'A': 'Action Group',
                'D': 'Dialogue Group',
                'SH': 'Scene Heading Group'
                }
        self.element2basestyle = {
                    'text':'Normal Text',
                    'h': 'Scene Heading',
                    'a': 'Action',
                    'c': 'Character',
                    'p': 'Parenthetical',
                    'd': 'Dialogue',
                    't': 'Transition',
                    's': 'Shot',
                    'uc': 'Unspoken Character'
                    }
        self.translate = Translate()
        self.read = Read()
        self.parse = Parse()
        
        # Scene heading element
    

class Read(object):

    def __init__(self):
        
        super(Read, self).__init__()

    def auto(self, filepath: str) -> pd.DataFrame:
        fp = Path(filepath)
        
        try:
            filepath = fp.resolve()
        except FileNotFoundError:
            print('file does not exist')
        
        # get file extension
        fextention = str(filepath).split('.')[-1]
        
        if fextention in ['txt', 'text']:
            dfsc = self.text(filepath)
            
        if fextention in ['doc', 'docx']:
            dfsc = self.docx(filepath)
            
        if fextention in ['xml']:
            dfsc = self.openxml(filepath)
        
        else:
            dfsc = None
            
        return dfsc
        
    @staticmethod
    def openxml(filepath: str,
                pat_sh=None,
                pat_d=None) -> pd.DataFrame:
        
        dfsc = pd.read_xml(filepath, xpath='paragraphs/para')
        dfsc['style'] = pd.read_xml(
            filepath, xpath='paragraphs/para/style')['basestyle']
        dfsc = dfsc.rename(columns={'style':'Type', 'text': 'Element'})
        
        # Define Columns
        dfsc['Grp'] = None
        dfsc['Scene'] = None
        
        # Assign Grp D and A
        dfsc.loc[dfsc['Type'].str.contains('Dialogue|Character'),
                 'Grp'] = 'D'
        dfsc.loc[dfsc['Type'].str.contains('Action'),
                 'Grp'] = 'A'
        
        # regenerate Scene Numbers
        idx_sh = dfsc[dfsc['Type']== 'Scene Heading'].index
        dfsc.loc[dfsc.index.isin(idx_sh), 'Grp'] = 'H'
        nscenes = len(idx_sh)
        dfsc.loc[dfsc.index.isin(idx_sh), 'Scene'] = list(range(1, nscenes+1))
        dfsc['Scene'].fillna(method='ffill', inplace=True)
        dfsc['Scene'].fillna(-1, inplace=True)
        dfsc['Scene'].astype('int')
       

        # Clean na Element
        dfsc.dropna(subset=['Element'], inplace=True)        
        dfsc = dfsc[['Scene', 'Element', 'Grp', 'Type']].copy()
        
        return dfsc
    
    @staticmethod
    def text(filepath: str,
             pat_sh=['INT./EXT.', 'INT/EXT', 'EXT', 'EXT\.', 'INT', 'INT\.', 
                      '内./外.', '内/外.', '内景', '外景', 
                      '内\.', '内,', '外\.', '外,',
                     ],
             pat_d=None) -> pd.DataFrame:
        """
        This functions opens screennplays in text, assumming it somewhat
        conforms to the screenplay format.

        Parameters
        ----------
        filepath : str
            DESCRIPTION.
        pat_sh : TYPE, optional
            DESCRIPTION. The default is None.
        pat_d : TYPE, optional
            DESCRIPTION. The default is None.

        Returns
        -------
        dfsc : TYPE
            DESCRIPTION.

        """
        
        # Read text file
        f = open(filepath)
        lines = f.readlines()
        f.close()
        
        dfsc = pd.DataFrame(lines)
        dfsc.columns = ['raw']
        
        # Define Columns
        dfsc['Grp'] = None
        dfsc['Type'] = None
        dfsc['Scene'] = None
        
        # identify Scene Headings
        idx_sh = dfsc[dfsc['raw'].str.contains('|'.join(pat_sh))].index
        dfsc.loc[dfsc.index.isin(idx_sh), 'Grp'] = 'H'
        dfsc.loc[dfsc.index.isin(idx_sh), 'Type'] = 'Scene Heading'
        
        # regenerate Scene Numbers
        nscenes = len(idx_sh)
        dfsc.loc[dfsc.index.isin(idx_sh), 'Scene'] = list(range(1, nscenes+1))
        dfsc['Scene'].fillna(method='ffill', inplace=True)
        dfsc['Scene'].fillna(-1, inplace=True)
        dfsc['Scene'].astype('int')
        # Remove leading and trailing spaces in Scene Heading
        dfsc.loc[dfsc['Grp'] == 'H', 'raw'] = \
            dfsc.loc[dfsc['Grp'] == 'H', 'raw'].apply(str.strip)
            
        # Identify Action
        dfsc['nspaces'] = dfsc['raw'].apply(lambda x: len(x)-len(x.lstrip()))
        mid = (dfsc['nspaces'].max() - dfsc['nspaces'].min()) //2 - 2
        dfsc.loc[(dfsc['Grp'] != 'H') & (dfsc['nspaces'] <= mid), 'Grp'] = 'A'
        dfsc.loc[dfsc['Grp'] == 'A', 'Type'] = 'Action'
        # Identify Dialogue
        dfsc.loc[(dfsc['Grp'] != 'H') & (dfsc['nspaces'] > mid), 'Grp'] = 'D'
        dfsc['raw'] = dfsc['raw'].apply(str.strip)
        
        # Identify Dchar in D Group
        idx_dchar = dfsc[(dfsc['Grp'] == 'D') & dfsc['raw'].str.isupper()].index
        dfsc.loc[dfsc.index.isin(idx_dchar), 'Type'] = 'Character'
        # Identify Dialog in D Group
        dfsc.loc[dfsc['Grp'] == 'D', 'Type'] = \
            dfsc.loc[dfsc['Grp'] == 'D', 'Type'].fillna('Dialogue')
            
        # Identify SHOT and Transition in A Group
        idx_shot_and_transition = dfsc[(dfsc['Grp'] == 'A') & dfsc['raw'].str.isupper()].index
        dfsc.loc[dfsc.index.isin(idx_shot_and_transition), 'Grp'] = 'S'
        dfsc.loc[dfsc.index.isin(idx_shot_and_transition), 'Type'] = 'Shot'
        
        pat_shot = ['FADE', 'CUT', 'DISSOLVE', 'INTERCUT']
        idx_transition = dfsc[(dfsc.index.isin(idx_shot_and_transition)) & 
                dfsc['raw'].str.contains('|'.join(pat_shot), flags=re.IGNORECASE)].index
        dfsc.loc[dfsc.index.isin(idx_transition), 'Grp'] = 'T'
        dfsc.loc[dfsc.index.isin(idx_transition), 'Type'] = 'Transition'
        
        # Combine Elements
        dfsc['nelement'] = None
        nelement = dfsc[dfsc['raw'] == ''].shape[0]
        dfsc.loc[(dfsc['raw'] == ''), 'nelement'] = list(range(nelement))
        dfsc['nelement'].fillna(method='bfill', inplace=True)
 
        # Group to Elements
        dfsc = dfsc.groupby(['Scene', 'nelement', 'Grp', 'Type'])['raw'].apply(
            lambda x: ' '.join(x)).reset_index()
        dfsc = dfsc[dfsc['raw'] != '']
        dfsc.rename(columns={'raw': 'Element'}, inplace=True)
 
        # Clean SC
        dfsc['Element'] = dfsc['Element'].apply(lambda x: re.sub(':SC:', '', x).strip())
        
        # Organize Return
        dfsc = dfsc[['Scene', 'Element', 'Grp', 'Type']].copy()
        
        return dfsc
    
    @staticmethod
    def docx(filepath: str, 
             pat_sh=['INT./EXT.', 'INT/EXT', 'EXT', 'EXT\.', 'INT', 'INT\.', 
                      '内./外.', '内/外.', '内景', '外景', 
                      '内\.', '内,', '外\.', '外,', '内 ', '外 ',
                     ],
             pat_shot=['FADE', 'CUT', 'DISSOLVE', 'INTERCUT'],
             pat_d=None) -> pd.DataFrame:
        '''
        This function reads word docx into a screenplay structured pd.DataFrame, 
        oneline per row.

        Parameters
        ----------
        filepath : str
            DESCRIPTION.
        sctype : str, optional
            DESCRIPTION. The default is None.

        Returns
        -------
        script : pd.DataFrame 
            DESCRIPTION.

        '''
        
        # Read Docx
        document = Document(filepath)
        num_p = len(document.paragraphs)
        dfsc = []
        for num in range(num_p):
            dfsc.append(document.paragraphs[num].text)
        dfsc = pd.DataFrame(dfsc)
        dfsc.columns = ['Element']

        # Define Columns
        dfsc['Scene'] = None
        dfsc['Grp'] = None
        dfsc['Type'] = None
        
        # identify Scene Headings
        idx_sh = dfsc[dfsc['Element'].str.contains('|'.join(pat_sh))].index
        dfsc.loc[dfsc.index.isin(idx_sh), 'Grp'] = 'H'
        dfsc.loc[dfsc.index.isin(idx_sh), 'Type'] = 'Scene Heading'
        
        # Regenerate Scene Numbers
        dfsc.loc[dfsc.index.isin(idx_sh), 'Scene'] = [i 
            for i in range(1, len(idx_sh)+1)]
        dfsc['Scene'].fillna(method='ffill', inplace=True)
        dfsc['Scene'].fillna(-1, inplace=True)
            
        # Identify Dialogue
        if not pat_d:
            pat_d = '[:：]'
        idx_d = dfsc[dfsc['Element'].str.contains(pat_d)].index
        dfsc.loc[idx_d, 'Grp'] = 'D'
        dfsc.loc[idx_d, 'Type'] = 'Dialogue'
        
        # Identify Action
        dfsc['Grp'] = dfsc['Grp'].fillna('A')
        dfsc['Type'] = dfsc['Type'].fillna('Action')
 
        # Identify SHOT in A Group
        idx_shot_and_transition = dfsc[(dfsc['Grp'] == 'A') & dfsc['raw'].str.isupper()].index
        dfsc.loc[dfsc.index.isin(idx_shot_and_transition), 'Grp'] = 'S'
        dfsc.loc[dfsc.index.isin(idx_shot_and_transition), 'Type'] = 'Shot'
        
        #Identify TRansition in A Group
        idx_transition = dfsc[(dfsc.index.isin(idx_shot_and_transition)) & 
                dfsc['raw'].str.contains('|'.join(pat_shot), flags=re.IGNORECASE)].index
        dfsc.loc[dfsc.index.isin(idx_transition), 'Grp'] = 'T'
        dfsc.loc[dfsc.index.isin(idx_transition), 'Type'] = 'Transition'
        
        # Clean Element with na
        dfsc['Element'] = dfsc['Element'].apply(lambda x: str(x).strip())
        dfsc = dfsc[dfsc['Element'] != ''].reset_index(drop=True)


        return dfsc
            
    @staticmethod
    def pdf(filepath: str,
                pat_sh=None,
                pat_d=None) -> pd.DataFrame:    
        pass

    @staticmethod
    def extract_location(x, 
                         pat_location:str ='[-——\.,]+'):
        if x['Element']:
            location = str(x['Element'])
        else: return ''
        if x['IE']:
            location = re.sub(re.escape(str(x['IE'])), '', location)
        if x['Time']:
            location = re.sub(re.escape(str(x['Time'])), '', location)
            
        location = re.sub(pat_location, '', location)
        return location.strip()




class Parse(object):
    
    def __init__(self):
        super(Parse, self).__init__()    
    
    @staticmethod
    def scene_heading(df: pd.DataFrame,
                 pat_sh=None,
                 pat_time:str = '[-——]+\s*(.*)',
                 pat_location:str = '[-——\.,]+'
                 ) -> pd.DataFrame:
        
        dfsc = df.copy()
        
        # Define Columns if not exist
        if 'IE' not in dfsc.columns:
            dfsc['IE'] = None
        if 'Location' not in dfsc.columns:
            dfsc['Location'] = None      
        if 'Time' not in dfsc.columns:
            dfsc['Time'] = None          
            
            
        # Break down Scene Headings
        idx_sh = dfsc[dfsc['Grp'] == 'H'].index
        
        # Extract Location
        if not pat_sh:
            pat_sh = ['INT./EXT.', 'INT/EXT', 'EXT', 'EXT\.', 'INT', 'INT\.', 
                      '内./外.', '内/外.', '内景', '外景', 
                      '内\.', '内,', '外\.', '外,',
                     ]
        dfsc.loc[dfsc.index.isin(idx_sh), 'IE'] = \
            dfsc.loc[dfsc.index.isin(idx_sh), 'Element'].str.extract(
                '({})'.format('|'.join(pat_sh)), expand=False)
        
        # Extract Time
        dfsc.loc[dfsc.index.isin(idx_sh), 'Time'] = \
            dfsc.loc[dfsc.index.isin(idx_sh), 'Element'].str.extract(
                pat_time, expand=False)

        
        # Extract Location
        dfsc.loc[dfsc.index.isin(idx_sh), 'Location'] = \
            dfsc.loc[dfsc.index.isin(idx_sh), :].apply(
                lambda x: Read.extract_location(x, pat_location=pat_location), 
                axis=1)
            
        return dfsc.copy()
    
class Elements(object):
    
    def __init__(self):
        super(Elements, self).__init__()
        
    def identify_scene_heading(self, sc, pattern: str = '\d*[.\s]') -> pd.DataFrame:
        sc.loc[sc['pcontent'].str.contains(pattern), 'ptype'] = 'h'
        return sc
    
    def identify_dialog(self, sc, pattern: str = u'[：:]') -> pd.DataFrame:
        sc.loc[sc['pcontent'].apply(lambda x: x[0:10]).str.contains(pattern), 'ptype'] = 'd'
        return sc
    
    def identify_action(self, sc, pattern: str = None) -> pd.DataFrame:
        if pattern:
            sc.loc[sc['pcontent'].str.contains(pattern), 'ptype'] = 'a'
        else:
            sc['ptype'].fillna(value='a', inplace=True)
        return sc
    
    # Scene Heading Subelements
    def identify_scene_number(self, 
                              sc, 
                              pattern: str, 
                              identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_scene_heading(sc, pattern)
        sc.loc[sc['ptype'] == 'h', 'h_number'] = sc.loc[sc['ptype'] == 'h', 'pcontent'].str.extract(pattern).iloc[:, 0]
        #sc['h_number'].fillna(method='ffill', inplace=True)
        return sc
    
    def identify_scene_inout(self, 
                              sc, 
                              pattern: str = r"[\w\S]*(.*?)[，,]", 
                              identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_scene_heading(sc, pattern)
        sc.loc[sc['ptype'] == 'h', 'h_inout'] = sc.loc[sc['ptype'] == 'h', 'pcontent'].str.extract(pattern).iloc[:, 0].str.lstrip(' ')
        #sc['h_inout'].fillna(method='ffill', inplace=True)
        return sc
        pass
        
    def identify_scene_title(self, 
                          sc, 
                          pattern: str = r'景[，,](\s*\w*.*)[，,]', 
                          identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_scene_heading(sc, pattern)
        sc.loc[sc['ptype'] == 'h', 'h_title'] = sc.loc[sc['ptype'] == 'h', 'pcontent'].str.extract(pattern).iloc[:, 0]
        #sc['h_title'].fillna(method='ffill', inplace=True)
        return sc

    def identify_scene_time(self, 
                          sc, 
                          pattern: str = r'[，,](\s*\w*.*$)', 
                          identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_scene_heading(sc, pattern)
        sc.loc[sc['ptype'] == 'h', 'h_time'] = sc.loc[sc['ptype'] == 'h', 'pcontent'].str.extract(pattern).iloc[:, 0]
        #sc['h_time'].fillna(method='ffill', inplace=True)
        return sc
    
    # Dialog Subelements
    def identify_dialog_character(self,
                                  sc,
                                  pattern: str = r'(^.*$)[:：]',
                                  identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_dialog(sc, pattern)
        sc.loc[sc['ptype'] == 'd', 'd_character'] = sc.loc[sc['ptype'] == 'd', 'pcontent'].str.extract(pattern).iloc[:, 0]
        sc  = self.identify_dialog_character_parenthesis(sc)
        #sc['h_time'].fillna(method='ffill', inplace=True)
        return sc

    def identify_dialog_character_parenthesis(self,
                                  sc: pd.DataFrame,
                                  pattern: str = r'[（(](.*)[）)]',
                                  identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_dialog(sc, pattern)
        sc.loc[sc['ptype'] == 'd', 'd_character_parenthesis'] = sc.loc[sc['ptype'] == 'd', 'd_character'].str.extract(pattern).iloc[:, 0]
        sc.loc[sc['ptype'] == 'd', 'd_character'] = sc.loc[sc['ptype'] == 'd', 'd_character'].str.replace(pattern, '')
        #sc['h_time'].fillna(method='ffill', inplace=True)
        return sc        
                        
    def identify_dialog_dialog(self,
                                 sc,
                                 pattern: str = r'[:：](.*)',
                                 identify_type_first: bool = False) -> pd.DataFrame:
       if identify_type_first:
           sc = self.identify_dialog(sc, pattern)
       sc.loc[sc['ptype'] == 'd', 'd_dialog'] = sc.loc[sc['ptype'] == 'd', 'pcontent'].str.extract(pattern).iloc[:, 0]
       sc.loc[sc['ptype'] == 'd', 'd_dialog'] = sc.loc[sc['ptype'] == 'd', 'd_dialog'].apply(lambda x: x.lstrip(' ').rstrip(' '))
       sc  = self.identify_dialog_parenthetisis(sc)
       return sc
   
    def identify_dialog_parenthetisis(self,
                                      sc: pd.DataFrame,
                                      pattern: str = r'[（(](.*)[）)]',
                                      identify_type_first: bool = False) -> pd.DataFrame:
       if identify_type_first:
           sc = self.identify_dialog(sc, pattern)
       sc.loc[sc['ptype'] == 'd', 'd_dialog_parenthesis'] = sc.loc[sc['ptype'] == 'd', 'd_dialog']. \
           str.extractall(r'[（(](.*?)[）)]').reset_index().groupby('level_0')[0].agg(';'.join).str.split(';')


       #sc['h_time'].fillna(method='ffill', inplace=True)
       return sc
   
######################################################
class Export(object):
    
    def __init__(self):
        super(Export, self).__init__()
    
    class to_OpenXML(object):
        
        def __init__(self):
            pass
    
        def to_openxml(self, sc: pd.DataFrame, save: bool = False, file_path: str  = None) -> str:
            '''
            
    
            Parameters
            ----------
            sc : pd.DataFrame
                DESCRIPTION.
            save : str, optional
                DESCRIPTION. The default is None.
            file_path : str, optional
                DESCRIPTION. The default is None.
    
            Returns
            -------
            str
                DESCRIPTION.
    
            '''
            sc['formatted'] = None
    
            header_xml = '<?xml version="1.0" encoding="UTF-8"?>\n'
            header_document = '<document type="Open Screenplay Format document" version="30">\n'
            document_info = '\t<info>\n\t</info>\n'            
            header_format = '\t<styles>\n\t</styles>\n'
            header_paragraphs = '\t<paragraphs>\n'
            
            footer_paragraphs = '\t</paragraphs>\n'
            section_titlepage = '\t<titlepage>\n\t</titlepage>\n'
            section_lists = '\t<lists>\n\t</lists>\n'
            footer_document = '</document>\n'
                
            sc = self.heading(sc)
            sc = self.action(sc)
            sc = self.dialog(sc)
    
            formatted_output = header_xml \
                             + header_document \
                             + document_info \
                             + header_format \
                             + header_paragraphs \
                             + sc['formatted'].sum() \
                             + footer_paragraphs \
                             + section_titlepage \
                             + section_lists \
                             + footer_document \
                                 
            if save:
                with codecs.open(file_path, "w", 'utf-8-sig') as f:
                    print(formatted_output, file=f)
                f.close()
                
            return formatted_output
        
        
        def heading(self, sc: pd.DataFrame) -> pd.DataFrame:
            sc.loc[sc['ptype'] == 'h', 'formatted'] = \
                    sc.loc[sc['ptype'] == 'h', 'h_inout'].apply(lambda x: str(x) + '. ') \
                +   sc.loc[sc['ptype'] == 'h', 'h_title'] \
                +   sc.loc[sc['ptype'] == 'h', 'h_time'].apply(lambda x: ' - ' + str(x))
            
            sc.loc[sc['ptype'] == 'h', 'formatted'] = sc.loc[sc['ptype'] == 'h', 'formatted'].apply(
                lambda x: '\t\t<para>\n\t\t\t<style basestylename="Scene Heading"/>\n\t\t\t<text>'
                        + str(x)
                        + '</text>\n\t\t</para>\n'            
                )
            return sc
        
        def action(self, sc: pd.DataFrame) -> pd.DataFrame:
            sc.loc[sc['ptype'] == 'a', 'formatted'] = \
                sc.loc[sc['ptype'] == 'a', 'pcontent'].apply(lambda x: str(x).lstrip('\s').rstrip('\s').replace('&', '&amp;'))
            sc.loc[sc['ptype'] == 'a', 'formatted'] = sc.loc[sc['ptype'] == 'a', 'formatted'].apply(
                lambda x: '\t\t<para>\n\t\t\t<style basestylename="Action"/>\n\t\t\t<text>'
                        + str(x)
                        + '</text>\n\t\t</para>\n'
                )
            return sc
    
        def dialog(self, sc: pd.DataFrame) -> pd.DataFrame:
            
            #
            sc.loc[sc['ptype'] == 'd', 'formatted'] = \
                  sc.loc[sc['ptype'] == 'd', 'd_character'] \
                + sc.loc[sc['ptype'] == 'd', 'd_character_parenthesis'].apply(lambda x: '(' + str(x) + ')' if not pd.isnull(x) else '')
            
            sc.loc[sc['ptype'] == 'd', 'formatted'] = sc.loc[sc['ptype'] == 'd', 'formatted'].apply(
                lambda x: '\t\t<para>\n\t\t\t<style basestylename="Character"/>\n\t\t\t<text>'
                        + str(x)
                        + '</text>\n\t\t</para>\n'
                )
                 
            sc.loc[(sc['ptype'] == 'd') & (sc['d_dialog_parenthesis'].isna()), 'formatted'] = \
                  sc.loc[(sc['ptype'] == 'd') & (sc['d_dialog_parenthesis'].isna()), 'formatted'] \
                + sc.loc[(sc['ptype'] == 'd') & (sc['d_dialog_parenthesis'].isna()), 'd_dialog'].apply(
                    lambda x: '\t\t<para>\n\t\t\t<style basestylename="Dialogue"/>\n\t\t\t<text>'
                        + str(x)
                        + '</text>\n\t\t</para>\n'
                    )
            
            sc.loc[(sc['ptype'] == 'd') & (~sc['d_dialog_parenthesis'].isna()), 'formatted'] = \
                  sc.loc[(sc['ptype'] == 'd') & (~sc['d_dialog_parenthesis'].isna()), 'formatted'] \
                + sc.loc[(sc['ptype'] == 'd') & (~sc['d_dialog_parenthesis'].isna())].agg(self.format_d_dialog, axis=1)
                            
            return sc
        
        def format_d_dialog(self, x: pd.DataFrame) -> str:
            '''
            formats dialog for XML output, to be used as a function for pandas agg or apply.
            
            Parameters
            ----------
            x : one row or column of pd.DataFrame depending on axis,
                to be used with the pandas agg or apply function
                DESCRIPTION.
    
            Returns
            -------
            str
                DESCRIPTION.
    
            '''
            pattern = r'[（(].*?[）)]'
            d_split = x['d_dialog'].split(pattern)
            f_dialog = ''
            for i, d in enumerate(d_split):
                ddialog = '\t\t<para>\n\t\t\t<style basestylename="Dialogue"/>\n\t\t\t<text>' \
                    + str(d) \
                    + '</text>\n\t\t</para>\n'
                if i == 0:
                    f_dialog = f_dialog + ddialog
                else:
                    dparenthetical = '\t\t<para>\n\t\t\t<style basestylename="Parenthetical"/>\n\t\t\t<text>' \
                        + str(x['d_dialog_parenthesis'][i-1]) \
                        + '</text>\n\t\t</para>\n' 
                    f_dialog = f_dialog + dparenthetical + ddialog
            return f_dialog

######
class Translate(object):
    
    def __init__(self, sc: str = None):   
        super(Translate, self).__init__()
    
    
    @staticmethod    
    def Baidu(s:str='apple', lang_from='zh', lang_to='en'):
        time.sleep(2)
        #path_translink = 'http://api.fanyi.baidu.com/api/trans/vip/translate'
        path_translink = 'https://fanyi-api.baidu.com/api/trans/vip/translate'
        httpClient = None
        
        appid = '20200218000385369'
        passcode = 'phecLsPI8EnkvjRHQ8HU'
        salt = randint(1e9, 9e9)
        q = s
        for_sign = appid + q + str(salt) + passcode
        sign = hashlib.md5(for_sign.encode()).hexdigest()
        
        link_query = (path_translink + '?'
                      + 'appid=' + appid +'&'
                      + 'q=' + urllib.parse.quote(q) + '&'
                      + 'from=' + lang_from + '&'
                      + 'to=' + lang_to + '&'
                      + 'salt=' + str(salt) + '&'
                      + 'sign=' + sign)
        
        try:
            httpClient = http.client.HTTPConnection('api.fanyi.baidu.com')
            httpClient.request('GET', link_query)
    
            # response是HTTPResponse对象
            response = httpClient.getresponse()
            result_all = response.read().decode("utf-8")
            result = json.loads(result_all)
            print(result['trans_result'][0]['dst'])
            return result['trans_result'][0]['dst']  # str
    
        except Exception as e:
            print(e)
        finally:
            if httpClient:
               httpClient.close() 